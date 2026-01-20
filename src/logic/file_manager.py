import os
import json
import hashlib
from typing import Optional, Dict, Any
import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import font as tkfont

# DOCX support
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.text import WD_COLOR_INDEX
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False


STYLE_TAG_PREFIX = "pw_fontstyle_"  # must match formatting.py
COLOR_FG_TAG_PREFIX = "color_fg_"
COLOR_BG_TAG_PREFIX = "color_bg_"


# Optional PDF export dependency
try:
    # Works with both PyFPDF (fpdf) and fpdf2
    from fpdf import FPDF
    HAS_FPDF = True
except Exception:
    HAS_FPDF = False

# Optional PDF import dependency (for opening PDFs as editable text)
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except Exception:
    HAS_PYMUPDF = False

try:
    from pypdf import PdfReader  # fallback
    HAS_PYPDF = True
except Exception:
    HAS_PYPDF = False


class FileManager:
    """File open/save manager.

    Supported:
      - .txt: plain text
      - .pwp: native PyWord Pro JSON (preserves formatting tags)
      - .docx: rich import/export using python-docx runs
      - .pdf: imports text only (not true PDF editing)

    This module only manipulates the Tk Text widget content + tags.
    """

    def __init__(self, editor, root):
        self.editor = editor
        self.root = root
        self.current_file_path = None

        # If the last opened document came from a PDF, we treat it as an
        # imported text document (not safe to overwrite the original PDF).
        self._opened_from_pdf = False

        # Keep references to dynamically created Tk Font objects used in tag
        # configurations so they don't get garbage collected.
        self._tag_font_refs = {}  # tag_name -> tkinter.font.Font

    # ----------------------------
    # Public API
    # ----------------------------

    def open_file(self):
        file_types = [
            ("Supported", "*.pwp *.txt *.docx *.pdf"),
            ("PyWord Pro", "*.pwp"),
            ("Text", "*.txt"),
            ("Word", "*.docx"),
            ("PDF", "*.pdf"),
            ("All", "*.*"),
        ]
        path = filedialog.askopenfilename(filetypes=file_types)
        if not path:
            return

        ext = os.path.splitext(path)[1].lower()
        try:
            if ext == ".pwp":
                self._open_pwp(path)
            elif ext == ".docx":
                if not HAS_DOCX:
                    raise RuntimeError("Opening .docx requires python-docx. Install: pip install python-docx")
                self._open_docx(path)
            elif ext == ".pdf":
                content = self._read_pdf_as_text(path)
                self._opened_from_pdf = True
                # Don't allow "Save" to overwrite a PDF with plain text.
                self.current_file_path = None
                self._set_editor_content(content)
                self.root.title(f"PyWord Pro - {os.path.basename(path)} (PDF Imported)")
            else:
                # default: treat as text
                with open(path, "r", encoding="utf-8") as f:
                    content = f.read()
                self._opened_from_pdf = False
                self.current_file_path = path
                self._set_editor_content(content)
                self.root.title(f"PyWord Pro - {os.path.basename(path)}")
        except Exception as e:
            messagebox.showerror("Error", f"Could not open file: {e}")

    def save_file(self):
        # If content came from a PDF import, we can't safely overwrite the
        # original PDF with plain text.
        if self._opened_from_pdf:
            messagebox.showinfo(
                "Save",
                "This document was imported from a PDF.\n\n"
                "Use 'Save As' to save as .pwp/.txt/.docx, or use 'Export PDF' to create a new PDF.",
            )
            self.save_as_file()
            return

        if self.current_file_path:
            self._write_file(self.current_file_path)
        else:
            self.save_as_file()

    def save_as_file(self):
        path = filedialog.asksaveasfilename(
            defaultextension=".pwp",
            filetypes=[
                ("PyWord Pro", "*.pwp"),
                ("Word", "*.docx"),
                ("Text", "*.txt"),
            ],
        )
        if not path:
            return

        if self._write_file(path):
            self.current_file_path = path
            self._opened_from_pdf = False
            self.root.title(f"PyWord Pro - {os.path.basename(path)}")

    def export_pdf(self):
        """Export the current editor text to a PDF file (plain text)."""
        if not HAS_FPDF:
            messagebox.showinfo(
                "PDF",
                "PDF export requires the 'fpdf' package.\n\nInstall it with: pip install fpdf",
            )
            return

        path = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF", "*.pdf")],
        )
        if not path:
            return

        try:
            content = self.editor.get("1.0", "end-1c")

            pdf = FPDF(format="A4", unit="mm")
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()

            # Built-in fonts (Helvetica/Times/Courier) are always available.
            # They are limited to latin-1; we replace unsupported chars to avoid crashes.
            pdf.set_font("Helvetica", size=12)

            for line in content.splitlines() or [""]:
                line = line.replace("\t", "    ")
                safe_line = line.encode("latin-1", "replace").decode("latin-1")
                pdf.multi_cell(0, 6, safe_line)

            pdf.output(path)
            messagebox.showinfo("PDF", f"Exported PDF successfully:\n{path}")
        except Exception as e:
            messagebox.showerror("PDF Export Error", str(e))

    # ----------------------------
    # Core write helpers
    # ----------------------------

    def _write_file(self, path: str) -> bool:
        try:
            ext = os.path.splitext(path)[1].lower()

            # Prevent corrupting a PDF by saving plain text to a .pdf path.
            if ext == ".pdf":
                messagebox.showinfo("Save", "To save as a PDF, use 'Export PDF' instead of 'Save'.")
                return False

            if ext == ".pwp":
                self._save_pwp(path)
                return True

            if ext == ".docx":
                if not HAS_DOCX:
                    raise RuntimeError("Saving .docx requires python-docx. Install: pip install python-docx")
                self._save_docx_runs(path)
                return True

            # default: .txt
            content = self.editor.get("1.0", "end-1c")
            with open(path, "w", encoding="utf-8") as f:
                f.write(content)
            return True
        except Exception as e:
            messagebox.showerror("Save Error", str(e))
            return False

    # ----------------------------
    # .PWP (native JSON) save/load
    # ----------------------------

    def _save_pwp(self, path: str) -> None:
        """Save text + formatting tags into a JSON file."""
        text = self.editor.get("1.0", "end-1c")

        # Collect tag ranges for all formatting tags.
        tags_payload = []
        tag_configs = {}

        for tname in self.editor.tag_names():
            if tname == "sel":
                continue
            ranges = self.editor.tag_ranges(tname)
            if not ranges:
                continue
            # store ranges
            rr = []
            for i in range(0, len(ranges), 2):
                rr.append([str(ranges[i]), str(ranges[i + 1])])
            tags_payload.append({"name": tname, "ranges": rr})

            # store tag config
            cfg = {}
            try:
                fg = self.editor.tag_cget(tname, "foreground")
                bg = self.editor.tag_cget(tname, "background")
                justify = self.editor.tag_cget(tname, "justify")
                fnt_name = self.editor.tag_cget(tname, "font")

                if fg:
                    cfg["foreground"] = fg
                if bg:
                    cfg["background"] = bg
                if justify:
                    cfg["justify"] = justify

                if fnt_name:
                    try:
                        fnt_obj = tkfont.nametofont(fnt_name)
                        cfg["font"] = {
                            "family": fnt_obj.cget("family"),
                            "size": int(fnt_obj.cget("size")),
                            "weight": fnt_obj.cget("weight"),
                            "slant": fnt_obj.cget("slant"),
                            "underline": int(fnt_obj.cget("underline")),
                            "overstrike": int(fnt_obj.cget("overstrike")),
                        }
                    except Exception:
                        pass
            except Exception:
                pass

            if cfg:
                tag_configs[tname] = cfg

        payload = {
            "version": 1,
            "text": text,
            "tags": tags_payload,
            "tag_configs": tag_configs,
        }

        with open(path, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)

    def _open_pwp(self, path: str) -> None:
        with open(path, "r", encoding="utf-8") as f:
            payload = json.load(f)

        if not isinstance(payload, dict) or payload.get("version") != 1:
            raise RuntimeError("Unsupported .pwp file version")

        self._tag_font_refs.clear()

        text = payload.get("text", "")
        self._set_editor_content(text)

        # Restore tag configurations first
        tag_configs = payload.get("tag_configs", {}) or {}
        for tname, cfg in tag_configs.items():
            try:
                font_info = cfg.get("font")
                if font_info:
                    # Recreate a Tk Font and keep a reference
                    f = tkfont.Font(
                        family=font_info.get("family") or "Calibri",
                        size=int(font_info.get("size") or 11),
                        weight=font_info.get("weight") or "normal",
                        slant=font_info.get("slant") or "roman",
                        underline=int(font_info.get("underline") or 0),
                        overstrike=int(font_info.get("overstrike") or 0),
                    )
                    self._tag_font_refs[tname] = f
                    self.editor.tag_configure(tname, font=f)

                if "foreground" in cfg:
                    self.editor.tag_configure(tname, foreground=cfg["foreground"])
                if "background" in cfg:
                    self.editor.tag_configure(tname, background=cfg["background"])
                if "justify" in cfg:
                    self.editor.tag_configure(tname, justify=cfg["justify"])
            except Exception:
                # Best-effort restore
                pass

        # Restore tag ranges
        for tag_rec in payload.get("tags", []) or []:
            try:
                tname = tag_rec.get("name")
                for a, b in tag_rec.get("ranges", []) or []:
                    self.editor.tag_add(tname, a, b)
            except Exception:
                pass

        self._opened_from_pdf = False
        self.current_file_path = path
        self.root.title(f"PyWord Pro - {os.path.basename(path)}")

    # ----------------------------
    # DOCX import/export (runs)
    # ----------------------------

    def _open_docx(self, path: str) -> None:
        """Import a DOCX and apply formatting tags."""
        doc = Document(path)

        self._tag_font_refs.clear()
        self.editor.delete("1.0", tk.END)

        first_paragraph = True
        for para in doc.paragraphs:
            if not first_paragraph:
                self.editor.insert("end", "\n")
            first_paragraph = False

            # paragraph alignment
            align_tag = None
            try:
                if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    align_tag = "center"
                elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                    align_tag = "right"
                else:
                    align_tag = "left"
            except Exception:
                align_tag = None

            # Insert runs
            for run in para.runs:
                txt = run.text or ""
                if not txt:
                    continue

                # IMPORTANT: Tk tag ranges are [start, end) (end-exclusive).
                # Using end-1c directly would miss the last character.
                start = self.editor.index("end-1c")
                self.editor.insert("end", txt)
                end_excl = self.editor.index("end-1c +1c")

                # Build style from run props
                family = None
                size_pt = None
                try:
                    family = run.font.name
                except Exception:
                    family = None

                try:
                    if run.font.size is not None:
                        size_pt = int(round(run.font.size.pt))
                except Exception:
                    size_pt = None

                bold = bool(run.bold) if run.bold is not None else False
                italic = bool(run.italic) if run.italic is not None else False
                underline = bool(run.underline) if run.underline is not None else False
                strike = bool(run.font.strike) if getattr(run.font, "strike", None) is not None else False

                # defaults if missing
                if family is None:
                    family = tkfont.Font(font=self.editor.cget("font")).cget("family")
                if size_pt is None:
                    size_pt = tkfont.Font(font=self.editor.cget("font")).cget("size")

                style_tag = self._ensure_style_tag(family, int(size_pt), bold, italic, underline, strike)
                self.editor.tag_add(style_tag, start, end_excl)

                # text color
                try:
                    rgb = run.font.color.rgb if run.font.color else None
                    if rgb:
                        hex_color = f"#{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x}"
                        fg_tag = f"{COLOR_FG_TAG_PREFIX}{hex_color}"
                        self.editor.tag_configure(fg_tag, foreground=hex_color)
                        self.editor.tag_add(fg_tag, start, end_excl)
                except Exception:
                    pass

                # highlight
                try:
                    hi = run.font.highlight_color
                    if hi is not None:
                        bg_hex = self._docx_highlight_to_hex(hi)
                        if bg_hex:
                            bg_tag = f"{COLOR_BG_TAG_PREFIX}{bg_hex}"
                            self.editor.tag_configure(bg_tag, background=bg_hex)
                            self.editor.tag_add(bg_tag, start, end_excl)
                except Exception:
                    pass

            # Apply alignment tag to the whole paragraph line (including newline)
            if align_tag:
                try:
                    # determine range for this paragraph line
                    line_no = int(self.editor.index("end-1c").split(".")[0])
                    ls = f"{line_no}.0"
                    le = f"{line_no}.0 lineend+1c"
                    self.editor.tag_configure(align_tag, justify=align_tag)
                    self.editor.tag_add(align_tag, ls, le)
                except Exception:
                    pass

        self._opened_from_pdf = False
        self.current_file_path = path
        self.root.title(f"PyWord Pro - {os.path.basename(path)}")

    def _save_docx_runs(self, path: str) -> None:
        """Export the editor content as a DOCX using runs to preserve formatting."""
        doc = Document()

        text = self.editor.get("1.0", "end-1c")
        lines = text.split("\n")

        for line_no, _line in enumerate(lines, start=1):
            # python-docx creates a new document with one empty paragraph.
            # Reuse that first paragraph for the first line to avoid a leading blank.
            if line_no == 1 and doc.paragraphs:
                p = doc.paragraphs[0]
            else:
                p = doc.add_paragraph()

            # alignment for this line
            align = self._get_alignment_for_line(line_no)
            if align == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == "right":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            line_start = f"{line_no}.0"
            line_end = f"{line_no}.0 lineend"

            segments = self._iter_run_segments(line_start, line_end)
            for seg in segments:
                seg_text = self.editor.get(seg["start"], seg["end"])
                if not seg_text:
                    continue

                run = p.add_run(seg_text)

                # Apply style from segment
                spec = seg.get("spec") or {}
                family = spec.get("family")
                size = spec.get("size")
                bold = spec.get("bold")
                italic = spec.get("italic")
                underline = spec.get("underline")
                strike = spec.get("strike")

                if family:
                    run.font.name = family
                if size:
                    run.font.size = Pt(int(size))
                run.bold = bool(bold)
                run.italic = bool(italic)
                run.underline = bool(underline)
                run.font.strike = bool(strike)

                # color
                fg = seg.get("fg")
                if fg:
                    try:
                        r, g, b = self._hex_to_rgb(fg)
                        run.font.color.rgb = RGBColor(r, g, b)
                    except Exception:
                        pass

                # highlight
                bg = seg.get("bg")
                if bg:
                    try:
                        run.font.highlight_color = self._hex_to_docx_highlight(bg)
                    except Exception:
                        # if mapping fails, ignore
                        pass

        doc.save(path)

    # ----------------------------
    # Tag helpers for DOCX runs
    # ----------------------------

    def _iter_run_segments(self, start: str, end: str):
        """Yield segments between start/end where run-affecting tags are constant."""
        boundaries = {self.editor.index(start), self.editor.index(end)}

        # Relevant tags: style tags + fg/bg color tags
        relevant = []
        for tname in self.editor.tag_names():
            if tname.startswith(STYLE_TAG_PREFIX) or tname.startswith(COLOR_FG_TAG_PREFIX) or tname.startswith(COLOR_BG_TAG_PREFIX):
                relevant.append(tname)

        for tname in relevant:
            try:
                ranges = self.editor.tag_ranges(tname)
            except tk.TclError:
                continue

            for i in range(0, len(ranges), 2):
                a = self.editor.index(ranges[i])
                b = self.editor.index(ranges[i + 1])

                # intersection with [start,end]
                if self.editor.compare(b, "<=", start) or self.editor.compare(a, ">=", end):
                    continue

                s = a if self.editor.compare(a, ">", start) else self.editor.index(start)
                e = b if self.editor.compare(b, "<", end) else self.editor.index(end)
                boundaries.add(s)
                boundaries.add(e)

        # sort boundaries
        def _key(ix: str):
            ix = self.editor.index(ix)
            ln, col = ix.split(".")
            return (int(ln), int(col))

        ordered = sorted(boundaries, key=_key)

        out = []
        for i in range(len(ordered) - 1):
            s = ordered[i]
            e = ordered[i + 1]
            try:
                if not self.editor.compare(s, "<", e):
                    continue
            except tk.TclError:
                continue

            tags = self.editor.tag_names(s)
            style_tag = next((t for t in tags if t.startswith(STYLE_TAG_PREFIX)), None)
            fg_tag = next((t for t in tags if t.startswith(COLOR_FG_TAG_PREFIX)), None)
            bg_tag = next((t for t in tags if t.startswith(COLOR_BG_TAG_PREFIX)), None)

            spec = self._style_spec_from_tag(style_tag)
            out.append({
                "start": s,
                "end": e,
                "spec": spec,
                "fg": self._color_from_tag(fg_tag, COLOR_FG_TAG_PREFIX),
                "bg": self._color_from_tag(bg_tag, COLOR_BG_TAG_PREFIX),
            })

        return out

    def _style_spec_from_tag(self, style_tag: Optional[str]) -> Dict[str, Any]:
        """Read effective font/style from a combined style tag."""
        # Defaults from editor base font
        base = tkfont.Font(font=self.editor.cget("font"))
        spec = {
            "family": base.cget("family"),
            "size": int(base.cget("size")),
            "bold": False,
            "italic": False,
            "underline": False,
            "strike": False,
        }

        if not style_tag:
            return spec

        # Prefer exporting the *base* font size (at 100% zoom) when available.
        # Our combined style tags include an `_s{size}_` marker.
        size_from_marker = False
        try:
            import re as _re
            m = _re.search(r"_s(\d+)_", style_tag)
            if m:
                spec["size"] = int(m.group(1))
                size_from_marker = True
        except Exception:
            pass

        # Flags from tag name (b1/i1/u1/o1)
        try:
            parts = style_tag[len(STYLE_TAG_PREFIX):].split("_")
            bits = {p[:1]: p[1:] for p in parts if len(p) == 2}
            spec["bold"] = bits.get("b") == "1"
            spec["italic"] = bits.get("i") == "1"
            spec["underline"] = bits.get("u") == "1"
            spec["strike"] = bits.get("o") == "1"
        except Exception:
            pass

        # Family/size from tag font (preferred)
        try:
            fnt_name = self.editor.tag_cget(style_tag, "font")
            if fnt_name:
                f = tkfont.nametofont(fnt_name)
                spec["family"] = f.cget("family")
                # Only override size if we didn't get a base size from the tag name.
                if not size_from_marker:
                    spec["size"] = int(f.cget("size"))
                # Keep consistency with flags
                spec["bold"] = spec["bold"] or (f.cget("weight") == "bold")
                spec["italic"] = spec["italic"] or (f.cget("slant") == "italic")
                spec["underline"] = spec["underline"] or bool(int(f.cget("underline")))
                spec["strike"] = spec["strike"] or bool(int(f.cget("overstrike")))
        except Exception:
            pass

        return spec

    def _color_from_tag(self, tag: Optional[str], prefix: str) -> Optional[str]:
        if not tag:
            return None
        if not tag.startswith(prefix):
            return None
        return tag[len(prefix):]

    def _get_alignment_for_line(self, line_no: int) -> Optional[str]:
        """Return left/center/right for a given line based on tags."""
        idx = f"{line_no}.0"
        try:
            tags = self.editor.tag_names(idx)
        except tk.TclError:
            return None
        for t in ("center", "right", "left"):
            if t in tags:
                return t
        return None

    # ----------------------------
    # Style tag creation
    # ----------------------------

    def _style_tag_name(self, family: str, size: int, b: bool, i: bool, u: bool, o: bool) -> str:
        safe = re_safe_family(family)
        fam_hash = hashlib.md5(family.encode("utf-8")).hexdigest()[:6]
        bits = (
            "b1" if b else "b0",
            "i1" if i else "i0",
            "u1" if u else "u0",
            "o1" if o else "o0",
        )
        return f"{STYLE_TAG_PREFIX}f{safe}{fam_hash}_s{int(size)}_" + "_".join(bits)

    def _ensure_style_tag(self, family: str, size: int, b: bool, i: bool, u: bool, o: bool) -> str:
        """Ensure a style tag exists and is configured."""
        tag = self._style_tag_name(family, size, b, i, u, o)
        if tag in self.editor.tag_names():
            return tag

        # Best-effort zoom scaling based on current editor base font vs 11.
        try:
            base_size = int(tkfont.Font(font=self.editor.cget("font")).cget("size"))
            zoom_ratio = base_size / 11.0
        except Exception:
            zoom_ratio = 1.0

        f = tkfont.Font(
            family=family,
            size=max(1, int(round(size * zoom_ratio))),
            weight="bold" if b else "normal",
            slant="italic" if i else "roman",
            underline=1 if u else 0,
            overstrike=1 if o else 0,
        )
        self._tag_font_refs[tag] = f
        self.editor.tag_configure(tag, font=f)
        return tag

    # ----------------------------
    # PDF import
    # ----------------------------

    def _read_pdf_as_text(self, path: str) -> str:
        """Extract readable text from a PDF.

        This is NOT full PDF editing. We import the PDF's text into the editor.
        Layout, images, and complex formatting won't be preserved.
        """
        # Preferred: PyMuPDF
        if HAS_PYMUPDF:
            try:
                doc = fitz.open(path)
                parts = []
                for page in doc:
                    parts.append(page.get_text("text"))
                doc.close()
                text = "\n\n".join([p.strip("\n") for p in parts]).strip()
                if text:
                    return text
            except Exception:
                pass

        # Fallback: pypdf
        if HAS_PYPDF:
            try:
                reader = PdfReader(path)
                parts = []
                for page in reader.pages:
                    parts.append(page.extract_text() or "")
                text = "\n\n".join([p.strip("\n") for p in parts]).strip()
                if text:
                    return text
            except Exception:
                pass

        # If we reached here, we couldn't extract any text.
        if not (HAS_PYMUPDF or HAS_PYPDF):
            raise RuntimeError(
                "PDF import requires 'pymupdf' (recommended) or 'pypdf'.\n"
                "Install one of them:\n  pip install pymupdf\n  pip install pypdf"
            )

        return (
            "[No extractable text found in this PDF]\n\n"
            "This often happens when the PDF is scanned images (not selectable text).\n"
            "If you need to edit a scanned PDF, you'd need OCR support."
        )

    # ----------------------------
    # Utilities
    # ----------------------------

    def _set_editor_content(self, content: str) -> None:
        self.editor.delete("1.0", tk.END)
        self.editor.insert("1.0", content)

    def _hex_to_rgb(self, hx: str):
        hx = hx.lstrip("#")
        return int(hx[0:2], 16), int(hx[2:4], 16), int(hx[4:6], 16)

    def _hex_to_docx_highlight(self, hx: str):
        """Map a hex color to the nearest built-in Word highlight."""
        # Word highlight is a limited palette.
        # We'll map a few common colors; default to yellow.
        hx = (hx or "").lower()
        if not hx.startswith("#") or len(hx) != 7:
            return WD_COLOR_INDEX.YELLOW

        mapping = {
            "#ffff00": WD_COLOR_INDEX.YELLOW,
            "#ff0": WD_COLOR_INDEX.YELLOW,
            "#00ff00": WD_COLOR_INDEX.BRIGHT_GREEN,
            "#00ffff": WD_COLOR_INDEX.TURQUOISE,
            "#ff00ff": WD_COLOR_INDEX.PINK,
            "#ff0000": WD_COLOR_INDEX.RED,
            "#0000ff": WD_COLOR_INDEX.BLUE,
            "#c0c0c0": WD_COLOR_INDEX.GRAY_25,
            "#808080": WD_COLOR_INDEX.GRAY_50,
        }
        return mapping.get(hx, WD_COLOR_INDEX.YELLOW)

    def _docx_highlight_to_hex(self, hi) -> Optional[str]:
        """Convert a python-docx highlight enum to an approximate hex color."""
        rev = {
            WD_COLOR_INDEX.YELLOW: "#ffff00",
            WD_COLOR_INDEX.BRIGHT_GREEN: "#00ff00",
            WD_COLOR_INDEX.TURQUOISE: "#00ffff",
            WD_COLOR_INDEX.PINK: "#ff00ff",
            WD_COLOR_INDEX.RED: "#ff0000",
            WD_COLOR_INDEX.BLUE: "#0000ff",
            WD_COLOR_INDEX.GRAY_25: "#c0c0c0",
            WD_COLOR_INDEX.GRAY_50: "#808080",
        }
        return rev.get(hi)


def re_safe_family(family: str) -> str:
    import re
    safe = re.sub(r"[^A-Za-z0-9]+", "-", family or "").strip("-")
    return safe or "font"
